import os
import sys
import json
import psycopg2
from psycopg2 import extras
import shutil
import urllib3
import uvicorn
import secrets
import re
import logging
import time
import io
import base64
import smtplib 
from email.mime.text import MIMEText 
from email.mime.multipart import MIMEMultipart
from datetime import datetime, timedelta
from typing import Optional, List, Dict, Any, Union

# --- KÜTÜPHANELER ---
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("❌ python-docx kütüphanesi eksik.")

# --- GÖRÜNTÜ İŞLEME KÜTÜPHANESİ ---
try:
    from PIL import Image
except ImportError:
    print("⚠️ Pillow (PIL) kütüphanesi eksik. Docker'da 'requirements.txt' içine 'pillow' eklediğinden emin ol.")

# FastAPI ve Yan Ürünleri
from fastapi import FastAPI, Depends, HTTPException, status, UploadFile, File, Form
from fastapi.responses import StreamingResponse
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, EmailStr
from passlib.context import CryptContext
from jose import JWTError, jwt
import requests

# PDF İşleme
try:
    from pypdf import PdfReader
except ImportError:
    pass

# --- LOGLAMA ---
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger("NexusKernel")

# --- AYARLAR ---
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
GROQ_KEY = os.getenv("GROQ_API_KEY")
TAVILY_KEY = os.getenv("TAVILY_API_KEY")
DATABASE_URL = os.getenv("DATABASE_URL")
SECRET_KEY = os.getenv("SECRET_KEY", "nexus_key_2026")

# --- MAIL AYARLARI ---
EMAIL_USER = os.getenv("EMAIL_USER") 
EMAIL_PASSWORD = os.getenv("EMAIL_PASSWORD") 
SMTP_SERVER = "smtp.gmail.com"
SMTP_PORT = 587

ALGORITHM = "HS256"
VECTOR_DB_PATH = "nexus_vector_db"

# --- AI ÇEKİRDEĞİ ---
embeddings = None
vector_db = None

try:
    from langchain_huggingface import HuggingFaceEmbeddings
    from langchain_chroma import Chroma
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    from tavily import TavilyClient
    
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2", model_kwargs={'device': 'cpu'})
    vector_db = Chroma(persist_directory=VECTOR_DB_PATH, embedding_function=embeddings)
    
    logger.info("✅ AI Çekirdeği ve Hafıza Aktif.")
except Exception as e:
    logger.warning(f"AI Bileşenleri Yüklenemedi (Opsiyonel): {e}")

# --- DB BAĞLANTISI ---
def get_db_connection():
    if not DATABASE_URL:
        # Fallback 
        return psycopg2.connect("postgresql://postgres:postgres@backend/nexusdb")
    return psycopg2.connect(DATABASE_URL)

def init_db():
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        
        # 1. Kullanıcılar
        cur.execute("""
            CREATE TABLE IF NOT EXISTS users (
                id SERIAL PRIMARY KEY, 
                email TEXT UNIQUE, 
                password_hash TEXT, 
                full_name TEXT, 
                firm_name TEXT, 
                created_at TEXT
            );
        """)
        
        # 2. Sohbet Geçmişi
        cur.execute("""
            CREATE TABLE IF NOT EXISTS history (
                id SERIAL PRIMARY KEY, 
                user_id INTEGER, 
                title TEXT, 
                content TEXT, 
                analysis TEXT, 
                dilekce_taslagi TEXT, 
                date TEXT
            );
        """)

        # 3. Davalarım Tablosu
        cur.execute("""
            CREATE TABLE IF NOT EXISTS cases (
                id SERIAL PRIMARY KEY,
                user_id INTEGER,
                case_name TEXT,
                court_name TEXT,
                status TEXT, 
                summary TEXT,
                risks TEXT,
                todo_list TEXT,
                created_at TEXT
            );
        """)
        
        # Sütun Kontrolleri (Migration)
        try:
            cur.execute("ALTER TABLE cases ADD COLUMN IF NOT EXISTS summary TEXT;")
            conn.commit()
        except: conn.rollback()
        
        try:
            cur.execute("ALTER TABLE cases ADD COLUMN IF NOT EXISTS risks TEXT;")
            conn.commit()
        except: conn.rollback()
        
        try:
            cur.execute("ALTER TABLE cases ADD COLUMN IF NOT EXISTS todo_list TEXT;")
            conn.commit()
        except: conn.rollback()

        try:
            # --- KRİTİK: OCR SÜTUNU ---
            cur.execute("ALTER TABLE cases ADD COLUMN IF NOT EXISTS ocr_text TEXT;")
            conn.commit()
        except: conn.rollback()

        # 4. Hatırlatıcılar
        cur.execute("""
            CREATE TABLE IF NOT EXISTS reminders (
                id SERIAL PRIMARY KEY,
                case_id INTEGER,
                description TEXT,
                due_date TEXT,
                is_completed BOOLEAN DEFAULT FALSE
            );
        """)

        # 5. Dava Dokümanları
        cur.execute("""
            CREATE TABLE IF NOT EXISTS case_documents (
                id SERIAL PRIMARY KEY,
                case_id INTEGER,
                file_name TEXT,
                uploaded_at TEXT
            );
        """)

        # 6. Finans
        cur.execute("""
            CREATE TABLE IF NOT EXISTS finance (
                id SERIAL PRIMARY KEY,
                case_id INTEGER,
                type TEXT, 
                amount NUMERIC,
                description TEXT,
                date TEXT
            );
        """)

        # 7. Kişiler
        cur.execute("""
            CREATE TABLE IF NOT EXISTS parties (
                id SERIAL PRIMARY KEY,
                case_id INTEGER,
                role TEXT,
                name TEXT,
                contact_info TEXT
            );
        """)

        # 8. Notlar
        cur.execute("""
            CREATE TABLE IF NOT EXISTS case_notes (
                id SERIAL PRIMARY KEY,
                case_id INTEGER,
                note TEXT,
                created_at TEXT
            );
        """)

        # 9. OTP Kodları
        cur.execute("""
            CREATE TABLE IF NOT EXISTS otp_codes (
                email TEXT PRIMARY KEY,
                code TEXT,
                expires_at TIMESTAMP
            );
        """)
        
        conn.commit()
        cur.close()
        conn.close()
        logger.info("✅ Veritabanı Tabloları Hazır.")
    except Exception as e:
        logger.error(f"DB Init Error: {e}")

init_db()

app = FastAPI(title="Nexus Ultimate Law AI")

app.add_middleware(
    CORSMiddleware, 
    allow_origins=["*"], 
    allow_credentials=True, 
    allow_methods=["*"], 
    allow_headers=["*"]
)

pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="auth/login")

# --- MODELLER ---
class UserRegister(BaseModel):
    email: EmailStr
    password: str
    full_name: str
    firm_name: Optional[str] = "Avukat"

class CaseCreate(BaseModel):
    case_name: str
    court_name: Optional[str] = "Belirtilmedi"

class CaseStatusUpdate(BaseModel):
    status: str

class FinanceCreate(BaseModel):
    type: str
    amount: float
    description: str

class PartyCreate(BaseModel):
    role: str
    name: str
    contact_info: str

class NoteCreate(BaseModel):
    note: str

class ForgotPasswordRequest(BaseModel):
    email: EmailStr

class VerifyOTPRequest(BaseModel):
    email: EmailStr
    code: str

# --- YARDIMCI FONKSİYONLAR ---
def get_password_hash(p): return pwd_context.hash(p)
def verify_password(p, h): return pwd_context.verify(p, h)
def create_token(data): return jwt.encode({**data, "exp": datetime.utcnow() + timedelta(days=30)}, SECRET_KEY, algorithm=ALGORITHM)

def clean_ai_json(text):
    """AI bazen JSON'ın başına sonuna yazı ekler, bunu temizler ve HATA TOLERANSI sağlar."""
    try:
        text = re.sub(r"```json", "", text)
        text = re.sub(r"```", "", text)
        match = re.search(r"\{.*\}", text, re.DOTALL)
        if match:
            return json.loads(match.group())
        return json.loads(text)
    except:
        # JSON Parsing Hatası Olursa Regex ile Kurtarmayı Dene
        try:
            logger.warning("⚠️ JSON Parsing Başarısız, Manuel Regex deneniyor...")
            summary_match = re.search(r'"summary":\s*"(.*?)"', text, re.DOTALL)
            full_text_match = re.search(r'"full_text":\s*"(.*?)"', text, re.DOTALL)
            risks_match = re.search(r'"risks":\s*"(.*?)"', text, re.DOTALL)
            todo_list_match = re.search(r'"todo_list":\s*"(.*?)"', text, re.DOTALL)
            
            return {
                "summary": summary_match.group(1) if summary_match else "Özet çıkarılamadı.",
                "full_text": full_text_match.group(1) if full_text_match else text[:2500],
                "risks": risks_match.group(1) if risks_match else "",
                "todo_list": todo_list_match.group(1) if todo_list_match else "",
                "reminders": []
            }
        except:
            return {"full_text": text[:2000], "summary": "Otomatik analiz başarısız."}

async def get_current_user(token: str = Depends(oauth2_scheme)):
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        email = payload.get("sub")
    except: raise HTTPException(401)
    
    conn = get_db_connection()
    cur = conn.cursor(cursor_factory=extras.DictCursor)
    cur.execute("SELECT * FROM users WHERE email = %s", (email,))
    user = cur.fetchone()
    cur.close()
    conn.close()
    if not user: raise HTTPException(401)
    return dict(user)

def classify_intent(query: str):
    keywords = ["merhaba", "selam", "günaydın", "test", "naber"]
    if len(query.split()) < 3 and query.lower().strip() in keywords:
        return False
    return True 

def send_otp_email(to_email: str, code: str):
    if not EMAIL_USER or not EMAIL_PASSWORD:
        logger.error("Email ayarları eksik!")
        return False

    subject = "Nexus Doğrulama Kodu"
    body = f"""
    Merhaba,
    
    Nexus Hukuk Sistemine giriş veya şifre sıfırlama talebiniz için doğrulama kodunuz:
    
    {code}
    
    Bu kod 5 dakika süreyle geçerlidir.
    
    Saygılarımızla,
    Nexus AI Team
    """

    msg = MIMEMultipart()
    msg['From'] = EMAIL_USER
    msg['To'] = to_email
    msg['Subject'] = subject
    msg.attach(MIMEText(body, 'plain'))

    try:
        server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT)
        server.starttls()
        server.login(EMAIL_USER, EMAIL_PASSWORD)
        server.sendmail(EMAIL_USER, to_email, msg.as_string())
        server.quit()
        logger.info(f"Mail gönderildi: {to_email}")
        return True
    except Exception as e:
        logger.error(f"Mail gönderme hatası: {e}")
        return False

# --- VISION ANALİZ MOTORU (GÜÇLENDİRİLMİŞ) ---
def analyze_image_with_ai(image_data):
    """
    Groq Llama 3.2 90B Vision Modelini kullanarak görüntüyü analiz eder.
    """
    if not GROQ_KEY:
        logger.error("Groq API Key eksik!")
        return None

    # Görüntüyü base64 formatına çevir
    base64_image = base64.b64encode(image_data).decode('utf-8')
    image_url = f"data:image/jpeg;base64,{base64_image}"

    # AI Promptu
    prompt = """
    Sen uzman bir Türk Hukuk Asistanısın. Sana verilen görüntüyü analiz et.
    
    GÖREVLERİN:
    1. OCR (TAM METİN): Görüntüdeki yazıları, el yazısı olsa bile kelimesi kelimesine oku ve 'full_text' alanına yaz.
    2. ANALİZ: Bu metinden şu bilgileri çıkar:
       - Davanın Özeti (Kısa ve net)
       - Riskler (Aleyhimize ne olabilir?)
       - Yapılacaklar (Avukat ne yapmalı?)
       - Kritik Tarihler (Varsa)
    
    SADECE ŞU JSON FORMATINDA CEVAP VER:
    {
        "full_text": "Resimde okunan tüm metin buraya (satırları koru)...",
        "summary": "Metnin özeti...",
        "risks": "- Risk 1\n- Risk 2",
        "todo_list": "- Görev 1\n- Görev 2",
        "reminders": [
            {"date": "15.05.2026", "description": "Ödeme günü"}
        ]
    }
    """

    try:
        # MODEL DEĞİŞİKLİĞİ: 90b (Çalışan Model)
        response = requests.post(
            "https://api.groq.com/openai/v1/chat/completions",
            headers={"Authorization": f"Bearer {GROQ_KEY}"},
            json={
                "model": "llama-3.2-90b-vision-preview", # <-- DÜZELTİLDİ: 11b YERİNE 90b
                "messages": [
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": prompt},
                            {"type": "image_url", "image_url": {"url": image_url}}
                        ]
                    }
                ],
                "response_format": {"type": "json_object"},
                "temperature": 0.1 # Daha tutarlı sonuç için
            }
        )
        
        # Loglama
        if response.status_code != 200:
            logger.error(f"Groq API Hatası: {response.text}")
            return None

        raw_content = response.json()['choices'][0]['message']['content']
        logger.info(f"AI Vision Başarılı: {raw_content[:100]}...")
        
        return clean_ai_json(raw_content)

    except Exception as e:
        logger.error(f"Vision API Error: {e}")
        return None

# ==========================================
# ENDPOINTLER
# ==========================================

@app.post("/auth/register")
def register(user: UserRegister):
    conn = get_db_connection(); cur = conn.cursor()
    try:
        cur.execute("INSERT INTO users (email, password_hash, full_name, firm_name, created_at) VALUES (%s, %s, %s, %s, %s)",
            (user.email, get_password_hash(user.password), user.full_name, user.firm_name, str(datetime.now())))
        conn.commit(); return {"status": "Kayıt başarılı"}
    except: raise HTTPException(400, "Email kullanımda")
    finally: cur.close(); conn.close()

@app.post("/auth/login")
def login(form: OAuth2PasswordRequestForm = Depends()):
    conn = get_db_connection(); cur = conn.cursor(cursor_factory=extras.DictCursor)
    cur.execute("SELECT * FROM users WHERE email = %s", (form.username,))
    user = cur.fetchone(); cur.close(); conn.close()
    if not user or not verify_password(form.password, user['password_hash']): raise HTTPException(401)
    return {"access_token": create_token({"sub": user['email']}), "token_type": "bearer"}

@app.post("/auth/forgot-password")
def forgot_password(request: ForgotPasswordRequest):
    conn = get_db_connection(); cur = conn.cursor()
    cur.execute("SELECT id FROM users WHERE email = %s", (request.email,))
    user = cur.fetchone()
    if not user:
        cur.close(); conn.close()
        raise HTTPException(status_code=404, detail="Bu e-posta adresiyle kayıtlı kullanıcı bulunamadı.")

    otp_code = "".join([str(secrets.randbelow(10)) for _ in range(6)])
    expires_at = datetime.now() + timedelta(minutes=5)

    try:
        cur.execute("""
            INSERT INTO otp_codes (email, code, expires_at) 
            VALUES (%s, %s, %s)
            ON CONFLICT (email) 
            DO UPDATE SET code = EXCLUDED.code, expires_at = EXCLUDED.expires_at;
        """, (request.email, otp_code, expires_at))
        
        conn.commit()
        send_otp_email(request.email, otp_code)
        
        return {"status": "ok", "message": "Doğrulama kodu gönderildi."}
        
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        cur.close(); conn.close()

@app.post("/auth/verify-otp")
def verify_otp(request: VerifyOTPRequest):
    conn = get_db_connection()
    cur = conn.cursor(cursor_factory=extras.DictCursor)
    
    try:
        cur.execute("SELECT * FROM otp_codes WHERE email = %s", (request.email,))
        record = cur.fetchone()
        
        if not record:
            raise HTTPException(status_code=400, detail="Kod talep edilmedi veya süresi doldu.")
        
        db_code = record['code']
        expiry = record['expires_at']
        
        if db_code != request.code:
            raise HTTPException(status_code=400, detail="Hatalı kod girdiniz.")
            
        if datetime.now() > expiry:
            raise HTTPException(status_code=400, detail="Kodun süresi dolmuş. Lütfen tekrar isteyin.")
            
        cur.execute("DELETE FROM otp_codes WHERE email = %s", (request.email,))
        conn.commit()
        return {"status": "ok", "message": "Kod doğrulandı."}
        
    except HTTPException as he:
        raise he
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        cur.close(); conn.close()

@app.post("/chat/upload")
async def upload_files_for_chat(files: List[UploadFile] = File(...), user: dict = Depends(get_current_user)):
    if not vector_db: raise HTTPException(500, "AI Hafızası Devre Dışı")
    
    processed_text = ""
    file_names = []
    
    for f in files:
        file_names.append(f.filename)
        try:
            if f.filename.endswith(".pdf"):
                reader = PdfReader(f.file)
                for page in reader.pages:
                    text = page.extract_text()
                    if text: processed_text += text + "\n"
            elif f.filename.endswith(".docx"):
                doc = Document(f.file)
                for para in doc.paragraphs:
                    processed_text += para.text + "\n"
        except Exception as e:
            logger.error(f"Dosya okuma hatası: {e}")

    if processed_text and len(processed_text) > 0:
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        chunks = text_splitter.split_text(str(processed_text))
        
        metadatas = [{"source": fname, "user_id": user['id'], "type": "chat_upload"} for fname in file_names for _ in chunks]
        vector_db.add_texts(texts=chunks, metadatas=metadatas)
        
        return {"status": "ok", "message": f"{len(files)} dosya işlendi."}
    
    return {"status": "error", "message": "Dosyalar okunamadı veya metin yok."}

@app.post("/chat/voice")
async def voice_to_text(file: UploadFile = File(...)):
    try:
        files_data = {"file": (file.filename, file.file, "audio/mpeg")}
        r = requests.post("https://api.groq.com/openai/v1/audio/transcriptions", headers={"Authorization": f"Bearer {GROQ_KEY}"}, files=files_data, data={"model": "whisper-large-v3"})
        return r.json()
    except Exception as e:
        return {"text": "", "error": str(e)}

@app.post("/cases/create")
def create_case(case: CaseCreate, user: dict = Depends(get_current_user)):
    conn = get_db_connection(); cur = conn.cursor()
    cur.execute("INSERT INTO cases (user_id, case_name, court_name, status, created_at) VALUES (%s, %s, %s, 'Aktif', %s) RETURNING id",
        (user['id'], case.case_name, case.court_name, str(datetime.now())))
    case_id = cur.fetchone()[0]
    conn.commit(); cur.close(); conn.close()
    return {"id": case_id, "message": "Dava dosyası açıldı."}

@app.get("/cases/list")
def list_cases(user: dict = Depends(get_current_user)):
    conn = get_db_connection(); cur = conn.cursor(cursor_factory=extras.DictCursor)
    cur.execute("SELECT * FROM cases WHERE user_id = %s ORDER BY id DESC", (user['id'],))
    rows = cur.fetchall(); cur.close(); conn.close()
    return [dict(r) for r in rows]

@app.delete("/cases/{case_id}")
def delete_case(case_id: int, user: dict = Depends(get_current_user)):
    conn = get_db_connection(); cur = conn.cursor()
    try:
        cur.execute("DELETE FROM reminders WHERE case_id = %s", (case_id,))
        cur.execute("DELETE FROM case_documents WHERE case_id = %s", (case_id,))
        cur.execute("DELETE FROM finance WHERE case_id = %s", (case_id,))
        cur.execute("DELETE FROM parties WHERE case_id = %s", (case_id,))
        cur.execute("DELETE FROM case_notes WHERE case_id = %s", (case_id,))
        cur.execute("DELETE FROM cases WHERE id = %s AND user_id = %s", (case_id, user['id']))
        conn.commit()
        return {"status": "ok", "message": "Dava dosyası ve tüm verileri silindi."}
    except Exception as e:
        conn.rollback()
        raise HTTPException(500, str(e))
    finally:
        cur.close(); conn.close()

@app.put("/cases/{case_id}/status")
def update_case_status(case_id: int, status_update: CaseStatusUpdate, user: dict = Depends(get_current_user)):
    conn = get_db_connection(); cur = conn.cursor()
    try:
        cur.execute("UPDATE cases SET status = %s WHERE id = %s AND user_id = %s", 
                    (status_update.status, case_id, user['id']))
        conn.commit()
        return {"status": "ok", "message": f"Dava durumu güncellendi: {status_update.status}"}
    except Exception as e:
        conn.rollback()
        raise HTTPException(500, str(e))
    finally:
        cur.close(); conn.close()

@app.post("/cases/{case_id}/upload")
async def upload_case_files(case_id: int, files: List[UploadFile] = File(...), user: dict = Depends(get_current_user)):
    conn = get_db_connection(); cur = conn.cursor(); full_text = ""
    for f in files:
        try:
            cur.execute("INSERT INTO case_documents (case_id, file_name, uploaded_at) VALUES (%s, %s, %s)", 
                        (case_id, f.filename, str(datetime.now())))
            if f.filename.endswith(".pdf"):
                reader = PdfReader(f.file)
                for page in reader.pages: 
                    text = page.extract_text()
                    if text: full_text += text + "\n"
            elif f.filename.endswith(".docx"):
                doc = Document(f.file)
                for para in doc.paragraphs: full_text += para.text + "\n"
        except Exception as e:
            logger.error(f"Dosya okuma hatası ({f.filename}): {e}")

    conn.commit() 

    if not full_text:
        cur.close(); conn.close()
        return {"status": "error", "msg": "Dosyalar okunamadı veya metin içermiyor."}

    prompt = f"Hukuk asistanısın. Metni analiz et (Özet, Riskler, Yapılacaklar, Tarihler):\n{full_text[:15000]}\nJSON VER: summary, risks, todo_list, reminders(date, description)"
    try:
        r = requests.post("https://api.groq.com/openai/v1/chat/completions",
            headers={"Authorization": f"Bearer {GROQ_KEY}"},
            json={
                "model": "llama-3.3-70b-versatile", 
                "messages": [{"role": "user", "content": prompt}], 
                "response_format": {"type": "json_object"}
            }
        )
        raw_content = r.json()['choices'][0]['message']['content']
        data = clean_ai_json(raw_content)
        
        summary = data.get("summary", "Özet çıkarılamadı.")
        risks = data.get("risks", "")
        todo = data.get("todo_list", "")
        reminders = data.get("reminders", [])
        
        cur.execute("UPDATE cases SET summary = %s, risks = %s, todo_list = %s WHERE id = %s", 
                    (summary, risks, todo, case_id))
        
        for rem in reminders:
            cur.execute("INSERT INTO reminders (case_id, description, due_date) VALUES (%s, %s, %s)",
                (case_id, rem.get('description', 'Hatırlatma'), rem.get('date', '')))
        
        conn.commit()
        cur.close(); conn.close()
        
        return {
            "status": "ok", 
            "processed_files": 1, 
            "summary": summary,
            "found_reminders": reminders
        }
    
    except Exception as e:
        cur.close(); conn.close()
        logger.error(f"AI Error: {e}")
        return {"status": "error", "msg": f"Analiz Hatası: {str(e)}"}

@app.post("/cases/{case_id}/upload-vision")
async def upload_case_image(case_id: int, file: UploadFile = File(...), user: dict = Depends(get_current_user)):
    """
    Görüntüyü analiz eder, OCR metnini 'ocr_text' sütununa ve 'Notlar' tablosuna kaydeder.
    """
    content = await file.read()
    conn = get_db_connection(); cur = conn.cursor()
    
    cur.execute("INSERT INTO case_documents (case_id, file_name, uploaded_at) VALUES (%s, %s, %s)", 
                (case_id, file.filename, str(datetime.now())))
    conn.commit()

    # AI Analizi
    ai_result = analyze_image_with_ai(content)
    
    if not ai_result:
        cur.close(); conn.close()
        # Eğer AI hata verirse (örneğin model kapalıysa), bunu kullanıcıya açıkça göster.
        return {"status": "error", "msg": "Görüntü analiz edilemedi. Logları kontrol edin veya model ismini güncelleyin."}

    full_text = ai_result.get("full_text", "")
    summary = ai_result.get("summary", "Özet çıkarılamadı.")
    risks = ai_result.get("risks", "")
    todo = ai_result.get("todo_list", "")
    
    try:
        cur.execute("UPDATE cases SET summary = %s, risks = %s, todo_list = %s, ocr_text = %s WHERE id = %s", 
                    (summary, risks, todo, full_text, case_id))
    except Exception as e:
        logger.error(f"DB Update Error (Sütun eksik olabilir mi?): {e}")
        conn.rollback()
        cur.execute("UPDATE cases SET summary = %s, risks = %s, todo_list = %s WHERE id = %s", 
                    (summary, risks, todo, case_id))
    
    if full_text:
        cur.execute("INSERT INTO case_notes (case_id, note, created_at) VALUES (%s, %s, %s)",
                    (case_id, f"📷 [VISION TARAMASI]:\n{full_text}", str(datetime.now().strftime("%d.%m.%Y %H:%M"))))

    for rem in ai_result.get("reminders", []):
        cur.execute("INSERT INTO reminders (case_id, description, due_date) VALUES (%s, %s, %s)",
            (case_id, rem.get('description', 'AI Hatırlatma'), rem.get('date', '')))
            
    conn.commit()
    cur.close(); conn.close()

    return {
        "status": "ok",
        "message": "Görüntü başarıyla analiz edildi.",
        "analysis": ai_result
    }

@app.get("/cases/{case_id}/details")
def get_case_details(case_id: int):
    conn = get_db_connection(); cur = conn.cursor(cursor_factory=extras.DictCursor)
    
    cur.execute("SELECT * FROM cases WHERE id = %s", (case_id,))
    case_row = cur.fetchone()
    if not case_row: return {"error": "Dava bulunamadı"}
    case = dict(case_row)
    
    data = {}
    for t in ["case_documents", "reminders", "finance", "parties", "case_notes"]:
        cur.execute(f"SELECT * FROM {t} WHERE case_id = %s ORDER BY id DESC", (case_id,))
        data[t.replace("case_", "")] = [dict(r) for r in cur.fetchall()]
        if t == "case_documents": data["documents"] = data.pop("documents") 

    cur.close(); conn.close()
    return {"case": case, **data}

@app.get("/calendar/global")
def get_global_calendar(user: dict = Depends(get_current_user)):
    conn = get_db_connection(); cur = conn.cursor(cursor_factory=extras.DictCursor)
    query = """
        SELECT r.*, c.case_name 
        FROM reminders r
        JOIN cases c ON r.case_id = c.id
        WHERE c.user_id = %s
        ORDER BY r.due_date ASC
    """
    cur.execute(query, (user['id'],))
    rows = cur.fetchall()
    cur.close(); conn.close()
    return [dict(r) for r in rows]

@app.post("/cases/{case_id}/finance")
def add_finance(case_id: int, item: FinanceCreate, user: dict = Depends(get_current_user)):
    conn = get_db_connection(); cur = conn.cursor()
    cur.execute("INSERT INTO finance (case_id, type, amount, description, date) VALUES (%s, %s, %s, %s, %s)",
                (case_id, item.type, item.amount, item.description, str(datetime.now().strftime("%d.%m.%Y"))))
    conn.commit(); cur.close(); conn.close()
    return {"status": "ok"}

@app.post("/cases/{case_id}/party")
def add_party(case_id: int, party: PartyCreate, user: dict = Depends(get_current_user)):
    conn = get_db_connection(); cur = conn.cursor()
    cur.execute("INSERT INTO parties (case_id, role, name, contact_info) VALUES (%s, %s, %s, %s)",
                (case_id, party.role, party.name, party.contact_info))
    conn.commit(); cur.close(); conn.close()
    return {"status": "ok"}

@app.post("/cases/{case_id}/note")
def add_note(case_id: int, note: NoteCreate, user: dict = Depends(get_current_user)):
    conn = get_db_connection(); cur = conn.cursor()
    cur.execute("INSERT INTO case_notes (case_id, note, created_at) VALUES (%s, %s, %s)",
                (case_id, note.note, str(datetime.now().strftime("%d.%m.%Y %H:%M"))))
    conn.commit(); cur.close(); conn.close()
    return {"status": "ok"}

@app.get("/sor")
async def ask(soru: str, user: dict = Depends(get_current_user)):
    if not classify_intent(soru):
        return {"cevap": f"Merhaba {user['full_name']}, belgelerinizi yükleyip analize başlayabiliriz.", "dilekce_taslagi": "", "riskler": []}

    context_text = ""
    if vector_db:
        try:
            docs = vector_db.similarity_search(soru, k=4) 
            context_text += "\n[YÜKLENEN BELGELERDEN BULUNANLAR]:\n" + "\n".join([d.page_content for d in docs])
        except: pass

    web_context = ""
    if TAVILY_KEY:
        try:
            tavily = TavilyClient(api_key=TAVILY_KEY)
            res = tavily.search(query=f"Yargıtay kararı {soru}", max_results=2)
            web_context = "\n[WEB/YARGITAY]:\n" + "\n".join([r['content'] for r in res['results']])
        except: pass

    prompt = f"""
    Sen Nexus Ultimate Hukuk Asistanısın.
    Soru: {soru}
    
    Yüklenen Belgeler: {context_text}
    Dış Bilgi: {web_context}
    
    GÖREVLER:
    1. Kullanıcı sorusunu detaylı yanıtla.
    2. Eğer dilekçe istenirse "dilekce_taslagi" alanını doldur.
    
    JSON Cevap Ver:
    {{
        "cevap": "Analiz metni...",
        "dilekce_taslagi": "...",
        "riskler": ["..."]
    }}
    """
    
    try:
        r = requests.post("https://api.groq.com/openai/v1/chat/completions",
            headers={"Authorization": f"Bearer {GROQ_KEY}"},
            json={"model": "llama-3.3-70b-versatile", "messages": [{"role": "user", "content": prompt}], "response_format": {"type": "json_object"}})
        
        data = clean_ai_json(r.json()['choices'][0]['message']['content'])
        
        conn = get_db_connection(); cur = conn.cursor()
        cur.execute("INSERT INTO history (user_id, title, content, analysis, dilekce_taslagi, date) VALUES (%s, %s, %s, %s, %s, %s)",
            (user['id'], soru[:30], soru, data.get("cevap"), data.get("dilekce_taslagi"), str(datetime.now())))
        conn.commit(); cur.close(); conn.close()
        
        return data
    except Exception as e:
        return {"cevap": f"Hata oluştu: {str(e)}", "riskler": []}

@app.get("/gecmis")
def history(user: dict = Depends(get_current_user)):
    conn = get_db_connection(); cur = conn.cursor(cursor_factory=extras.DictCursor)
    cur.execute("SELECT * FROM history WHERE user_id = %s ORDER BY id DESC LIMIT 50", (user['id'],))
    rows = cur.fetchall(); cur.close(); conn.close()
    return [dict(r) for r in rows]

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000) 